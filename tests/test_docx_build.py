"""真实 DOCX 发现与无 OCR 本地构建契约。"""

from __future__ import annotations

import hashlib
import io
from pathlib import Path

from docx import Document
from PIL import Image

from rag_app.chunking import (
    Chunker,
    ChunkerConfig,
    Utf8TokenCounter,
)
from rag_app.clients.model_services import EmbeddingResult
from rag_app.clients.resilience import (
    ExternalRequestRejectedError,
    ExternalServiceUnavailableError,
)
from rag_app.contracts import DocumentMetadata
from rag_app.index.build import (
    DocxBuildConfig,
    DocxBuildServices,
    DocxChunkBuilder,
    OcrElementProcessor,
    discover_docx_sources,
)
from rag_app.index.qdrant import IndexedChunk
from rag_app.ocr.models import DEFAULT_OCR_REVISION, OcrLine, OcrResponse
from rag_app.parsers import DocxParser
from rag_app.retrieval.bm25 import QdrantBm25Encoder
from rag_app.state import StateStore
from rag_app.state.models import SourceVersion, VersionState

_DEFAULT_METADATA = DocumentMetadata(
    document_status="active",
    authority_level="official",
    effective_from=None,
    effective_to=None,
)


class _DeterministicEmbedder:
    """只验证批次顺序的确定性测试替身。"""

    def embed(
        self,
        texts: tuple[str, ...],
        *,
        instruction: str,
    ) -> EmbeddingResult:
        del instruction
        return EmbeddingResult(
            vectors=tuple((float(len(text)), 0.0, 1.0) for text in texts),
            calls=(),
        )


class _OcrClient:
    """返回固定合成文本的 OCR 接缝替身。"""

    def __init__(self) -> None:
        self.calls = 0

    def recognize(
        self,
        media_bytes: bytes,
        *,
        media_type: str,
        media_sha256: str,
    ) -> OcrResponse:
        self.calls += 1
        assert media_type == "image/png"
        assert hashlib.sha256(media_bytes).hexdigest() == media_sha256
        return OcrResponse(
            media_sha256=media_sha256,
            ocr_revision=DEFAULT_OCR_REVISION,
            text="公开合成 OCR 文本",
            confidence=0.96,
            lines=(
                OcrLine(
                    text="公开合成 OCR 文本",
                    confidence=0.96,
                    bbox=(0, 0, 20, 20),
                ),
            ),
            width=20,
            height=20,
            elapsed_ms=2,
        )


class _FlakyOcrClient(_OcrClient):
    """首次不可用、后续恢复的 OCR 接缝替身。"""

    def recognize(
        self,
        media_bytes: bytes,
        *,
        media_type: str,
        media_sha256: str,
    ) -> OcrResponse:
        if self.calls == 0:
            self.calls += 1
            raise ExternalServiceUnavailableError("OCR 暂时不可用。")
        return super().recognize(
            media_bytes,
            media_type=media_type,
            media_sha256=media_sha256,
        )


class _RejectedOcrClient(_OcrClient):
    """始终返回终态请求拒绝的 OCR 接缝替身。"""

    def recognize(
        self,
        media_bytes: bytes,
        *,
        media_type: str,
        media_sha256: str,
    ) -> OcrResponse:
        del media_bytes, media_type, media_sha256
        self.calls += 1
        raise ExternalRequestRejectedError("OCR 请求被拒绝。")


def _write_synthetic_docx(path: Path, *, image_count: int) -> None:
    image_stream = io.BytesIO()
    Image.new("RGB", (20, 20), "white").save(
        image_stream,
        format="PNG",
    )
    document = Document()
    document.add_heading("公开合成标题", level=1)
    document.add_paragraph("公开合成正文")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "公开值"
    for _ in range(image_count):
        document.add_picture(io.BytesIO(image_stream.getvalue()))
    document.save(str(path))


def test_discover_frozen_docx_inputs(tmp_path: Path) -> None:
    docs_dir = tmp_path / "input"
    docs_dir.mkdir()
    nested = docs_dir / "nested"
    nested.mkdir()
    _write_synthetic_docx(docs_dir / "alpha.docx", image_count=0)
    _write_synthetic_docx(nested / "beta.DOCX", image_count=0)
    (docs_dir / "ignored.docx:Zone.Identifier").write_text(
        "[ZoneTransfer]",
        encoding="utf-8",
    )

    discovered = discover_docx_sources(docs_dir)

    assert [item.source_path for item in discovered] == [
        "alpha.docx",
        "nested/beta.DOCX",
    ]
    assert all("Zone.Identifier" not in item.source_path for item in discovered)


def test_build_all_docx_without_running_ocr(tmp_path: Path) -> None:
    docs_dir = tmp_path / "input"
    docs_dir.mkdir()
    _write_synthetic_docx(docs_dir / "synthetic.docx", image_count=2)
    state = StateStore(tmp_path / "state.sqlite3")
    state.initialize()
    pipeline_fingerprint = f"sha256:{'1' * 64}"
    chunker = Chunker(
        ChunkerConfig(
            target_tokens=384,
            hard_max_tokens=512,
            overlap_tokens=64,
        ),
        Utf8TokenCounter(),
        pipeline_fingerprint=pipeline_fingerprint,
    )
    discovered_sources = discover_docx_sources(docs_dir)
    builder = DocxChunkBuilder(
        config=DocxBuildConfig(
            input_root=docs_dir,
            ocr_revision="unselected",
            embedding_instruction="测试契约，不调用外部模型",
            metadata_by_source={
                source.source_path: _DEFAULT_METADATA
                for source in discovered_sources
            },
        ),
        services=DocxBuildServices(
            parser=DocxParser(),
            chunker=chunker,
            embedder=_DeterministicEmbedder(),
            sparse_encoder=QdrantBm25Encoder(
                tokenizer="multilingual",
                language="none",
            ),
            state=state,
        ),
    )

    indexed: list[IndexedChunk] = []
    for discovered in discovered_sources:
        version = SourceVersion(
            source_id=f"src_{discovered.content_sha256[:32]}",
            doc_version=f"sha256:{discovered.content_sha256}",
            content_sha256=discovered.content_sha256,
            source_path=discovered.source_path,
            pipeline_fingerprint=pipeline_fingerprint,
            state=VersionState.STAGING,
            job_id="job_test",
            chunk_count=None,
            error_code=None,
        )
        indexed.extend(builder(discovered.source_path, version))

    assert indexed
    assert state.count_ocr_results(
        ocr_revision="unselected",
        state="pending",
    ) == 1
    assert state.count_ocr_results(
        ocr_revision="unselected",
        error_code="GPU_OCR_PENDING_SELECTION",
    ) == 1
    assert state.count_media_references(
        ocr_revision="unselected",
        state="pending",
    ) == 2
    assert state.count_media_references(
        ocr_revision="unselected",
        media_type="image/png",
    ) == 2
    assert all(item.chunk.contains_ocr is False for item in indexed)
    assert all(len(item.dense) == 3 for item in indexed)


def test_builder_calls_ocr_once_and_indexes_image_evidence(
    tmp_path: Path,
) -> None:
    input_root = tmp_path / "input"
    input_root.mkdir()
    image_stream = io.BytesIO()
    Image.new("RGB", (20, 20), "white").save(
        image_stream,
        format="PNG",
    )
    document = Document()
    document.add_paragraph("公开合成正文")
    document.add_picture(io.BytesIO(image_stream.getvalue()))
    document_path = input_root / "synthetic.docx"
    document.save(str(document_path))
    discovered = discover_docx_sources(input_root)[0]
    state = StateStore(tmp_path / "state.sqlite3")
    state.initialize()
    fingerprint = "sha256:" + ("2" * 64)
    ocr_client = _OcrClient()
    builder = DocxChunkBuilder(
        config=DocxBuildConfig(
            input_root=input_root,
            ocr_revision=DEFAULT_OCR_REVISION,
            embedding_instruction="公开合成测试",
            metadata_by_source={
                discovered.source_path: _DEFAULT_METADATA
            },
            minimum_ocr_confidence=0.80,
        ),
        services=DocxBuildServices(
            parser=DocxParser(),
            chunker=Chunker(
                ChunkerConfig(128, 256, 32),
                Utf8TokenCounter(),
                pipeline_fingerprint=fingerprint,
            ),
            embedder=_DeterministicEmbedder(),
            sparse_encoder=QdrantBm25Encoder(
                tokenizer="multilingual",
                language="none",
            ),
            state=state,
            ocr_client=ocr_client,
        ),
    )
    version = SourceVersion(
        source_id=f"src_{discovered.content_sha256[:32]}",
        doc_version=f"sha256:{discovered.content_sha256}",
        content_sha256=discovered.content_sha256,
        source_path=discovered.source_path,
        pipeline_fingerprint=fingerprint,
        state=VersionState.STAGING,
        job_id="job_synthetic",
        chunk_count=None,
        error_code=None,
    )

    chunks = builder(discovered.source_path, version)
    cached_chunks = builder(discovered.source_path, version)

    assert ocr_client.calls == 1
    assert state.count_ocr_results(
        ocr_revision=DEFAULT_OCR_REVISION,
        state="succeeded",
    ) == 1
    assert any(
        item.chunk.contains_ocr
        and item.chunk.text == "公开合成 OCR 文本"
        for item in chunks
    )
    assert cached_chunks == chunks


def test_public_ocr_processor_enriches_image_once(
    tmp_path: Path,
) -> None:
    input_root = tmp_path / "input"
    input_root.mkdir()
    _write_synthetic_docx(input_root / "synthetic.docx", image_count=1)
    discovered = discover_docx_sources(input_root)[0]
    state = StateStore(tmp_path / "state.sqlite3")
    state.initialize()
    ocr_client = _OcrClient()
    version = SourceVersion(
        source_id=f"src_{discovered.content_sha256[:32]}",
        doc_version=f"sha256:{discovered.content_sha256}",
        content_sha256=discovered.content_sha256,
        source_path=discovered.source_path,
        pipeline_fingerprint="sha256:" + ("5" * 64),
        state=VersionState.STAGING,
        job_id="job_processor",
        chunk_count=None,
        error_code=None,
    )
    elements = DocxParser().parse(
        input_root / discovered.source_path,
        display_path=discovered.source_path,
    )
    processor = OcrElementProcessor(
        state=state,
        ocr_client=ocr_client,
        ocr_revision=DEFAULT_OCR_REVISION,
        minimum_confidence=0.80,
    )

    processed = processor.process(elements, version)
    cached = processor.process(elements, version)

    assert ocr_client.calls == 1
    assert [element.text for element in processed if element.binary_data] == [
        "公开合成 OCR 文本"
    ]
    assert cached == processed
    assert state.count_ocr_results(
        ocr_revision=DEFAULT_OCR_REVISION,
        state="succeeded",
    ) == 1


def test_builder_retries_transient_ocr_failure(
    tmp_path: Path,
) -> None:
    input_root = tmp_path / "input"
    input_root.mkdir()
    _write_synthetic_docx(input_root / "synthetic.docx", image_count=1)
    discovered = discover_docx_sources(input_root)[0]
    state = StateStore(tmp_path / "state.sqlite3")
    state.initialize()
    fingerprint = "sha256:" + ("3" * 64)
    ocr_client = _FlakyOcrClient()
    builder = DocxChunkBuilder(
        config=DocxBuildConfig(
            input_root=input_root,
            ocr_revision=DEFAULT_OCR_REVISION,
            embedding_instruction="公开合成测试",
            metadata_by_source={
                discovered.source_path: _DEFAULT_METADATA
            },
        ),
        services=DocxBuildServices(
            parser=DocxParser(),
            chunker=Chunker(
                ChunkerConfig(128, 256, 32),
                Utf8TokenCounter(),
                pipeline_fingerprint=fingerprint,
            ),
            embedder=_DeterministicEmbedder(),
            sparse_encoder=QdrantBm25Encoder(
                tokenizer="multilingual",
                language="none",
            ),
            state=state,
            ocr_client=ocr_client,
        ),
    )
    version = SourceVersion(
        source_id=f"src_{discovered.content_sha256[:32]}",
        doc_version=f"sha256:{discovered.content_sha256}",
        content_sha256=discovered.content_sha256,
        source_path=discovered.source_path,
        pipeline_fingerprint=fingerprint,
        state=VersionState.STAGING,
        job_id="job_retry",
        chunk_count=None,
        error_code=None,
    )

    first = builder(discovered.source_path, version)
    second = builder(discovered.source_path, version)

    assert not any(item.chunk.contains_ocr for item in first)
    assert any(item.chunk.contains_ocr for item in second)
    assert ocr_client.calls == 2
    assert state.count_ocr_results(
        ocr_revision=DEFAULT_OCR_REVISION,
        state="succeeded",
    ) == 1
    assert state.count_ocr_results(
        ocr_revision=DEFAULT_OCR_REVISION,
        error_code="OCR_SERVICE_UNAVAILABLE",
    ) == 0


def test_builder_keeps_terminal_ocr_rejection_cached(
    tmp_path: Path,
) -> None:
    input_root = tmp_path / "input"
    input_root.mkdir()
    _write_synthetic_docx(input_root / "synthetic.docx", image_count=1)
    discovered = discover_docx_sources(input_root)[0]
    state = StateStore(tmp_path / "state.sqlite3")
    state.initialize()
    fingerprint = "sha256:" + ("4" * 64)
    ocr_client = _RejectedOcrClient()
    builder = DocxChunkBuilder(
        config=DocxBuildConfig(
            input_root=input_root,
            ocr_revision=DEFAULT_OCR_REVISION,
            embedding_instruction="公开合成测试",
            metadata_by_source={
                discovered.source_path: _DEFAULT_METADATA
            },
        ),
        services=DocxBuildServices(
            parser=DocxParser(),
            chunker=Chunker(
                ChunkerConfig(128, 256, 32),
                Utf8TokenCounter(),
                pipeline_fingerprint=fingerprint,
            ),
            embedder=_DeterministicEmbedder(),
            sparse_encoder=QdrantBm25Encoder(
                tokenizer="multilingual",
                language="none",
            ),
            state=state,
            ocr_client=ocr_client,
        ),
    )
    version = SourceVersion(
        source_id=f"src_{discovered.content_sha256[:32]}",
        doc_version=f"sha256:{discovered.content_sha256}",
        content_sha256=discovered.content_sha256,
        source_path=discovered.source_path,
        pipeline_fingerprint=fingerprint,
        state=VersionState.STAGING,
        job_id="job_rejected",
        chunk_count=None,
        error_code=None,
    )

    first = builder(discovered.source_path, version)
    second = builder(discovered.source_path, version)

    assert not any(item.chunk.contains_ocr for item in first)
    assert not any(item.chunk.contains_ocr for item in second)
    assert ocr_client.calls == 1
    assert state.count_ocr_results(
        ocr_revision=DEFAULT_OCR_REVISION,
        error_code="OCR_REQUEST_REJECTED",
    ) == 1
