import zipfile
from collections.abc import Callable
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches
from lxml import etree
from PIL import Image

from rag_app.chunking import Chunker, ChunkerConfig, Utf8TokenCounter
from rag_app.contracts import DocumentMetadata, ElementKind, OcrState
from rag_app.parsers.docx import DocxParser, DocxParserLimits, UnsafeDocxError

_WORD_NAMESPACE = (
    "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
)
_UNSUPPORTED_NAMESPACE = "urn:rag-test:unsupported"
_PIPELINE_FINGERPRINT = "sha256:" + "f" * 64
_SOURCE_ID = "src_" + "1" * 32
_DOC_VERSION = "sha256:" + "a" * 64
_METADATA = DocumentMetadata(
    document_status="active",
    authority_level="official",
    effective_from=None,
    effective_to=None,
)


def _write_test_image(path: Path) -> None:
    image = Image.new("RGB", (80, 40), color="white")
    image.save(path, format="PNG")


def _rewrite_document_xml(
    path: Path,
    mutate: Callable[[etree._Element], None],
) -> None:
    rewritten = path.with_name(f"{path.stem}-rewritten.docx")
    with (
        zipfile.ZipFile(path) as source,
        zipfile.ZipFile(rewritten, "w") as target,
    ):
        for item in source.infolist():
            content = source.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(content)
                mutate(root)
                content = etree.tostring(
                    root,
                    xml_declaration=True,
                    encoding="UTF-8",
                    standalone=True,
                )
            target.writestr(item, content)
    rewritten.replace(path)


def _body(root: etree._Element) -> etree._Element:
    body = root.find(f"{{{_WORD_NAMESPACE}}}body")
    assert body is not None
    return body


def _wrap_paragraph(
    root: etree._Element,
    *,
    paragraph_index: int,
    wrapper: etree._Element,
) -> None:
    body = _body(root)
    paragraphs = [
        child
        for child in body
        if etree.QName(child).localname == "p"
    ]
    paragraph = paragraphs[paragraph_index]
    position = body.index(paragraph)
    body.remove(paragraph)
    wrapper.append(paragraph)
    body.insert(position, wrapper)


def test_docx_parser_preserves_order_and_locators(tmp_path: Path) -> None:
    image_path = tmp_path / "evidence.png"
    _write_test_image(image_path)
    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("总则", level=1)
    document.add_paragraph("正文证据")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "值"
    table.cell(1, 0).text = "状态"
    table.cell(1, 1).text = "有效"
    document.add_picture(str(image_path), width=Inches(1))
    document.save(docx_path)

    elements = DocxParser().parse(docx_path, display_path="资料/sample.docx")

    assert [element.kind for element in elements] == [
        ElementKind.HEADING,
        ElementKind.PARAGRAPH,
        ElementKind.TABLE,
        ElementKind.IMAGE,
    ]
    assert elements[1].locator.heading_path == ("总则",)
    assert elements[1].locator.paragraph_index == 1
    assert elements[2].text == "字段 | 值\n状态 | 有效"
    assert elements[2].locator.table_index == 1
    assert elements[3].locator.image_index == 1
    assert elements[3].media_type == "image/png"
    assert elements[3].ocr_state == OcrState.PENDING
    assert elements[3].binary_data is not None


def test_docx_parser_rejects_zip_bomb_metadata(tmp_path: Path) -> None:
    docx_path = tmp_path / "oversized.docx"
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "x" * 5000)

    parser = DocxParser(
        DocxParserLimits(
            max_file_bytes=20_000,
            max_uncompressed_bytes=1000,
            max_entry_bytes=10_000,
            max_entries=10,
            max_compression_ratio=20,
            timeout_seconds=5.0,
        )
    )

    with pytest.raises(UnsafeDocxError, match="解压总量"):
        parser.parse(docx_path, display_path="oversized.docx")


def test_docx_parser_rejects_path_traversal(tmp_path: Path) -> None:
    docx_path = tmp_path / "traversal.docx"
    with zipfile.ZipFile(docx_path, "w") as archive:
        archive.writestr("../outside", "bad")

    with pytest.raises(UnsafeDocxError, match="非法归档路径"):
        DocxParser().parse(docx_path, display_path="traversal.docx")


def test_zone_identifier_is_not_a_docx_input(tmp_path: Path) -> None:
    marker = tmp_path / "sample.docx:Zone.Identifier"
    marker.write_text("[ZoneTransfer]", encoding="utf-8")

    with pytest.raises(UnsafeDocxError, match=r"Zone\.Identifier"):
        DocxParser().parse(marker, display_path=marker.name)


def test_docx_parser_marks_list_item_without_changing_paragraph_kind(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "list.docx"
    document = Document()
    document.add_paragraph("列表证据", style="List Bullet")
    document.save(docx_path)

    elements = DocxParser().parse(docx_path, display_path=docx_path.name)

    assert len(elements) == 1
    assert elements[0].kind == ElementKind.PARAGRAPH
    assert elements[0].list_level == 0


def test_duplicate_headings_are_stable_unique_and_rename_safe(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "duplicate-headings.docx"
    document = Document()
    for _ in range(2):
        document.add_heading("重复标题", level=1)
        document.add_paragraph("相同正文")
    document.save(docx_path)
    parser = DocxParser()

    first = parser.parse(docx_path, display_path="旧名称.docx")
    repeated = parser.parse(docx_path, display_path="旧名称.docx")
    renamed = parser.parse(docx_path, display_path="新名称.docx")

    assert first == repeated
    assert [item.locator.heading_index for item in first] == [1, 1, 2, 2]
    assert len({item.element_id for item in first}) == len(first)
    assert [item.element_id for item in first] == [
        item.element_id for item in renamed
    ]


def test_paragraph_preserves_tab_and_explicit_line_break(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "controls.docx"
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("甲")
    run.add_tab()
    run.add_text("乙")
    run.add_break()
    run.add_text("丙")
    document.save(docx_path)

    elements = DocxParser().parse(docx_path, display_path=docx_path.name)

    assert [item.text for item in elements] == ["甲\t乙\n丙"]


def test_control_only_paragraph_before_heading_produces_no_evidence(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "blank-controls.docx"
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_tab()
    run.add_break()
    document.add_heading("有效标题", level=1)
    document.save(docx_path)

    elements = DocxParser().parse(docx_path, display_path=docx_path.name)
    chunks = Chunker(
        ChunkerConfig(
            target_tokens=100,
            hard_max_tokens=100,
            overlap_tokens=10,
        ),
        Utf8TokenCounter(),
        pipeline_fingerprint=_PIPELINE_FINGERPRINT,
    ).chunk(
        _SOURCE_ID,
        _DOC_VERSION,
        elements,
        metadata=_METADATA,
    )

    assert [element.text for element in elements] == ["有效标题"]
    assert all(chunk.text.strip() for chunk in chunks)
    assert all(chunk.embedding_text.strip() for chunk in chunks)


def test_ordinary_content_control_is_expanded_in_document_order(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "ordinary-control.docx"
    document = Document()
    document.add_paragraph("前")
    document.add_paragraph("受控正文")
    document.add_paragraph("后")
    document.save(docx_path)

    def wrap(root: etree._Element) -> None:
        control = etree.Element(f"{{{_WORD_NAMESPACE}}}sdt")
        etree.SubElement(control, f"{{{_WORD_NAMESPACE}}}sdtPr")
        content = etree.SubElement(
            control,
            f"{{{_WORD_NAMESPACE}}}sdtContent",
        )
        _wrap_paragraph(
            root,
            paragraph_index=1,
            wrapper=content,
        )
        body = _body(root)
        content_position = body.index(content)
        body.remove(content)
        control.append(content)
        body.insert(content_position, control)

    _rewrite_document_xml(docx_path, wrap)

    elements, audit = DocxParser().parse_with_audit(
        docx_path,
        display_path=docx_path.name,
    )

    assert [element.text for element in elements] == ["前", "受控正文", "后"]
    assert audit.ordinary_controls_parsed == 1


def test_table_of_contents_control_is_skipped_and_audited(
    tmp_path: Path,
) -> None:
    docx_path = tmp_path / "toc-control.docx"
    document = Document()
    document.add_paragraph("目录生成内容")
    document.add_paragraph("正文")
    document.save(docx_path)

    def wrap(root: etree._Element) -> None:
        control = etree.Element(f"{{{_WORD_NAMESPACE}}}sdt")
        properties = etree.SubElement(
            control,
            f"{{{_WORD_NAMESPACE}}}sdtPr",
        )
        doc_part = etree.SubElement(
            properties,
            f"{{{_WORD_NAMESPACE}}}docPartObj",
        )
        gallery = etree.SubElement(
            doc_part,
            f"{{{_WORD_NAMESPACE}}}docPartGallery",
        )
        gallery.set(
            f"{{{_WORD_NAMESPACE}}}val",
            "Table of Contents",
        )
        content = etree.SubElement(
            control,
            f"{{{_WORD_NAMESPACE}}}sdtContent",
        )
        _wrap_paragraph(
            root,
            paragraph_index=0,
            wrapper=content,
        )
        body = _body(root)
        content_position = body.index(content)
        body.remove(content)
        control.append(content)
        body.insert(content_position, control)

    _rewrite_document_xml(docx_path, wrap)

    elements, audit = DocxParser().parse_with_audit(
        docx_path,
        display_path=docx_path.name,
    )

    assert [element.text for element in elements] == ["正文"]
    assert audit.toc_controls_skipped == 1


@pytest.mark.parametrize("evidence_kind", ("text", "image"))
def test_unknown_body_child_with_evidence_fails_closed(
    tmp_path: Path,
    evidence_kind: str,
) -> None:
    docx_path = tmp_path / f"unknown-{evidence_kind}.docx"
    document = Document()
    if evidence_kind == "text":
        document.add_paragraph("未知结构正文")
    else:
        image_path = tmp_path / "unknown.png"
        _write_test_image(image_path)
        document.add_picture(str(image_path), width=Inches(1))
    document.save(docx_path)

    def wrap(root: etree._Element) -> None:
        wrapper = etree.Element(f"{{{_UNSUPPORTED_NAMESPACE}}}block")
        _wrap_paragraph(
            root,
            paragraph_index=0,
            wrapper=wrapper,
        )

    _rewrite_document_xml(docx_path, wrap)

    with pytest.raises(UnsafeDocxError, match=r"不支持.*证据"):
        DocxParser().parse(docx_path, display_path=docx_path.name)


def test_table_and_nested_table_images_keep_duplicate_references(
    tmp_path: Path,
) -> None:
    image_path = tmp_path / "shared.png"
    _write_test_image(image_path)
    docx_path = tmp_path / "table-images.docx"
    document = Document()
    document.add_picture(str(image_path), width=Inches(1))
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "外层"
    table.cell(0, 0).paragraphs[0].add_run().add_picture(
        str(image_path),
        width=Inches(1),
    )
    nested = table.cell(0, 1).add_table(rows=1, cols=1)
    nested.cell(0, 0).paragraphs[0].add_run().add_picture(
        str(image_path),
        width=Inches(1),
    )
    document.save(docx_path)

    first = DocxParser().parse(docx_path, display_path=docx_path.name)
    second = DocxParser().parse(docx_path, display_path=docx_path.name)
    images = [item for item in first if item.kind == ElementKind.IMAGE]

    assert first == second
    assert len(images) == 3
    assert [item.locator.image_index for item in images] == [1, 2, 3]
    assert len({item.content_sha256 for item in images}) == 1
    assert len({item.element_id for item in images}) == 3
