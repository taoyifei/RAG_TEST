"""为公开源码测试生成不含私有语料的评测数据与 DOCX。"""

from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from PIL import Image

from evaluation.dataset import EvaluationDataset

_CASE_COUNT = 60
_TUNING_CASE_COUNT = 45
_OCR_CASE_COUNT = 5
_UNANSWERABLE_INDEXES = frozenset({10, 20, 30, 40, 50})
_REQUIRED_CATEGORIES = (
    "ordinary",
    "numeric",
    "table",
    "ocr",
    "cross_chunk",
    "rewrite",
    "multiturn",
    "conflict",
    "unanswerable",
    "prompt_injection",
)


def synthetic_evaluation_dataset() -> EvaluationDataset:
    """创建满足生产 schema 和覆盖约束的 60 题合成数据。

    Args:
        无参数。

    Returns:
        含 45 道 tuning、15 道 holdout 和 5 道 OCR 题的数据集。

    """
    cases = []
    for index in range(1, _CASE_COUNT + 1):
        category = _REQUIRED_CATEGORIES[
            (index - 1) % len(_REQUIRED_CATEGORIES)
        ]
        is_ocr = index <= _OCR_CASE_COUNT
        is_answerable = index not in _UNANSWERABLE_INDEXES
        categories = ("ocr", category) if is_ocr else (category,)
        evidence = (
            {
                "document": "synthetic",
                "locator_contains": "图片1" if is_ocr else "段落1",
                "quote": None if is_ocr else "公开合成证据",
            },
        ) if is_answerable else ()
        cases.append(
            {
                "id": f"Q{index:03d}",
                "split": (
                    "tuning" if index <= _TUNING_CASE_COUNT else "holdout"
                ),
                "categories": categories,
                "history_questions": (
                    ("上一轮公开合成问题",)
                    if "multiturn" in categories
                    else ()
                ),
                "question": f"公开合成问题 {index}",
                "validation_state": "verified_text",
                "expected": {
                    "answerable": is_answerable,
                    "required_facts": (
                        ("公开合成事实",) if is_answerable else ()
                    ),
                    "refusal_code": (
                        None if is_answerable else "NO_EVIDENCE"
                    ),
                    "evidence": evidence,
                },
            }
        )
    return EvaluationDataset.model_validate(
        {
            "dataset_version": "synthetic-public-v1",
            "review_method": "人工设计的公开合成测试数据",
            "documents": {"synthetic": "synthetic.docx"},
            "cases": cases,
        }
    )


def write_synthetic_dataset(path: Path) -> EvaluationDataset:
    """把公开合成评测数据写入测试临时目录。

    Args:
        path: 临时 JSON 输出路径。

    Returns:
        与写入内容相同的已校验数据集。

    """
    dataset = synthetic_evaluation_dataset()
    path.write_text(
        dataset.model_dump_json(indent=2),
        encoding="utf-8",
    )
    return dataset


def write_synthetic_evidence_docx(path: Path) -> None:
    """创建同时含标题、正文和 PNG 的公开合成 DOCX。

    Args:
        path: 临时 DOCX 输出路径。

    Returns:
        无返回值。

    """
    image_stream = io.BytesIO()
    Image.new("RGB", (24, 24), "white").save(image_stream, format="PNG")
    document = Document()
    document.add_heading("公开合成标题", level=1)
    document.add_paragraph("公开合成证据")
    document.add_picture(io.BytesIO(image_stream.getvalue()))
    document.save(str(path))
