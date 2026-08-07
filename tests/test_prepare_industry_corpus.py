import hashlib
import json
import os
import sys
import zipfile
from collections.abc import Callable
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from rag_app.contracts import ElementKind
from rag_app.parsers.docx import DocxParser
from scripts.industry_corpus import (
    EXPECTED_INVENTORY,
    CorpusPreparationError,
    SourceSpec,
    clean_docx,
    heading_candidate,
    prepare_industry_corpus,
)
from scripts.industry_corpus.ooxml import OoxmlPreparationError

_GIT_SHA = "a" * 40
_SOURCE_DATE_EPOCH = 1_700_000_000
_RELATIONSHIP_NAMESPACE = (
    "http://schemas.openxmlformats.org/package/2006/relationships"
)
_WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _write_fake_soffice(path: Path) -> Path:
    script = f"""#!{sys.executable}
import os
import sys
import time
from pathlib import Path
from docx import Document

if '--version' in sys.argv:
    print('LibreOffice 24.2.0.0')
    raise SystemExit(0)
mode = os.environ.get('FAKE_SOFFICE_MODE', 'ok')
if mode == 'timeout':
    time.sleep(30)
if mode == 'nonzero':
    raise SystemExit(7)
output = Path(sys.argv[sys.argv.index('--outdir') + 1])
source = Path(sys.argv[-1])
if mode == 'non_docx':
    (output / f'{{source.stem}}.odt').write_bytes(b'not-docx')
    raise SystemExit(0)
document = Document()
document.add_paragraph('1. 目的')
document.add_paragraph('合成测试正文')
document.save(output / f'{{source.stem}}.docx')
if mode == 'extra':
    document.save(output / 'extra.docx')
if os.environ.get('FAKE_FAIL_STEM') == source.stem:
    (output / f'{{source.stem}}.docx').write_bytes(b'broken')
"""
    path.write_text(script, encoding="utf-8")
    path.chmod(0o700)
    return path


def _write_sources(root: Path, specs: tuple[SourceSpec, ...]) -> None:
    root.mkdir()
    for index, spec in enumerate(specs, start=1):
        name = spec.canonical_name
        if index % 2 == 0:
            name = name.replace(" ", "", 1)
        (root / name).write_bytes(f"legacy-{index}".encode())


def _create_docx(path: Path, paragraphs: tuple[str, ...]) -> None:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def _rewrite_archive(
    path: Path,
    mutate: Callable[[dict[str, bytes]], None],
) -> None:
    with zipfile.ZipFile(path) as source:
        parts = {
            item.filename: source.read(item)
            for item in source.infolist()
            if not item.is_dir()
        }
    mutate(parts)
    rewritten = path.with_name(f"{path.stem}-rewritten.docx")
    with zipfile.ZipFile(rewritten, "w", zipfile.ZIP_DEFLATED) as target:
        for name, payload in parts.items():
            target.writestr(name, payload)
    rewritten.replace(path)


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _prepare(  # noqa: PLR0913
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    *,
    specs: tuple[SourceSpec, ...],
    mode: str = "ok",
    output_name: str = "output",
    timeout_seconds: float = 2.0,
) -> Path:
    source = tmp_path / f"source-{output_name}"
    _write_sources(source, specs)
    converter = _write_fake_soffice(tmp_path / "soffice")
    monkeypatch.setenv("FAKE_SOFFICE_MODE", mode)
    result = prepare_industry_corpus(
        source_dir=source,
        output_root=tmp_path / output_name,
        libreoffice_path=converter,
        source_date_epoch=_SOURCE_DATE_EPOCH,
        generated_from_git_sha=_GIT_SHA,
        timeout_seconds=timeout_seconds,
        expected_inventory=specs,
    )
    return result.root


def test_missing_source_directory_fails_before_conversion(
    tmp_path: Path,
) -> None:
    converter = _write_fake_soffice(tmp_path / "soffice")

    with pytest.raises(CorpusPreparationError, match="source dir"):
        prepare_industry_corpus(
            source_dir=tmp_path / "missing",
            output_root=tmp_path / "output",
            libreoffice_path=converter,
            source_date_epoch=_SOURCE_DATE_EPOCH,
            generated_from_git_sha=_GIT_SHA,
            expected_inventory=(SourceSpec("GM-01 测试.doc"),),
        )


@pytest.mark.parametrize("unsafe_kind", ["symlink", "directory", "fifo"])
def test_source_symlink_and_special_files_are_rejected(
    tmp_path: Path,
    unsafe_kind: str,
) -> None:
    source = tmp_path / "source"
    source.mkdir()
    target = source / "GM-01 测试.doc"
    if unsafe_kind == "symlink":
        backing = tmp_path / "backing.doc"
        backing.write_bytes(b"legacy")
        target.symlink_to(backing)
    elif unsafe_kind == "directory":
        target.mkdir()
    else:
        os.mkfifo(target)
    converter = _write_fake_soffice(tmp_path / "soffice")

    with pytest.raises(CorpusPreparationError, match=r"symlink|特殊"):
        prepare_industry_corpus(
            source_dir=source,
            output_root=tmp_path / "output",
            libreoffice_path=converter,
            source_date_epoch=_SOURCE_DATE_EPOCH,
            generated_from_git_sha=_GIT_SHA,
            expected_inventory=(SourceSpec("GM-01 测试.doc"),),
        )


@pytest.mark.parametrize("inventory_case", ["missing", "extra", "duplicate"])
def test_source_inventory_must_be_exact(
    tmp_path: Path,
    inventory_case: str,
) -> None:
    source = tmp_path / "source"
    source.mkdir()
    if inventory_case != "missing":
        (source / "GM-01 测试.doc").write_bytes(b"one")
    if inventory_case == "extra":
        (source / "GM-02 额外.doc").write_bytes(b"two")
    if inventory_case == "duplicate":
        (source / "GM-01测试.doc").write_bytes(b"duplicate")
    converter = _write_fake_soffice(tmp_path / "soffice")

    with pytest.raises(CorpusPreparationError, match=r"MISMATCH|重复"):
        prepare_industry_corpus(
            source_dir=source,
            output_root=tmp_path / "output",
            libreoffice_path=converter,
            source_date_epoch=_SOURCE_DATE_EPOCH,
            generated_from_git_sha=_GIT_SHA,
            expected_inventory=(SourceSpec("GM-01 测试.doc"),),
        )


@pytest.mark.parametrize(
    ("mode", "error_code"),
    [("timeout", "LIBREOFFICE_TIMEOUT"), ("nonzero", "NONZERO_EXIT")],
)
def test_libreoffice_timeout_and_nonzero_are_stable_failures(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    mode: str,
    error_code: str,
) -> None:
    source = tmp_path / "source"
    specs = (SourceSpec("GM-01 测试.doc"),)
    _write_sources(source, specs)
    converter = _write_fake_soffice(tmp_path / "soffice")
    monkeypatch.setenv("FAKE_SOFFICE_MODE", mode)

    with pytest.raises(CorpusPreparationError, match=error_code):
        prepare_industry_corpus(
            source_dir=source,
            output_root=tmp_path / "output",
            libreoffice_path=converter,
            source_date_epoch=_SOURCE_DATE_EPOCH,
            generated_from_git_sha=_GIT_SHA,
            timeout_seconds=0.1,
            expected_inventory=specs,
        )


@pytest.mark.parametrize("mode", ["extra", "non_docx"])
def test_libreoffice_output_must_be_one_same_stem_docx(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    mode: str,
) -> None:
    source = tmp_path / "source"
    specs = (SourceSpec("GM-01 测试.doc"),)
    _write_sources(source, specs)
    converter = _write_fake_soffice(tmp_path / "soffice")
    monkeypatch.setenv("FAKE_SOFFICE_MODE", mode)

    with pytest.raises(CorpusPreparationError, match="OUTPUT_SET_INVALID"):
        prepare_industry_corpus(
            source_dir=source,
            output_root=tmp_path / "output",
            libreoffice_path=converter,
            source_date_epoch=_SOURCE_DATE_EPOCH,
            generated_from_git_sha=_GIT_SHA,
            expected_inventory=specs,
        )


@pytest.mark.parametrize(
    "dangerous_part",
    [
        "word/vbaProject.bin",
        "word/embeddings/object1.bin",
        "word/activeX/activeX1.bin",
    ],
)
def test_macro_ole_and_activex_parts_are_rejected(
    tmp_path: Path,
    dangerous_part: str,
) -> None:
    source = tmp_path / "unsafe.docx"
    _create_docx(source, ("正文",))
    _rewrite_archive(source, lambda parts: parts.update({dangerous_part: b"x"}))

    with pytest.raises(OoxmlPreparationError, match=r"macro|OLE|ActiveX"):
        clean_docx(
            source=source,
            destination=tmp_path / "clean.docx",
            canonical_name="GM-01 测试.docx",
            source_date_epoch=_SOURCE_DATE_EPOCH,
        )


def test_gm03_private_character_cleaning_is_idempotent(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    _create_docx(source, (f"制度{chr(0xE004)}正文",))

    first_audit = clean_docx(
        source=source,
        destination=first,
        canonical_name="GM-03 质量管理制度.docx",
        source_date_epoch=_SOURCE_DATE_EPOCH,
    )
    second_audit = clean_docx(
        source=first,
        destination=second,
        canonical_name="GM-03 质量管理制度.docx",
        source_date_epoch=_SOURCE_DATE_EPOCH,
    )

    assert first_audit.removed_private_character_count == 1
    assert second_audit.removed_private_character_count == 0
    assert first_audit.visible_text_sha256 == second_audit.visible_text_sha256
    assert _sha256(first) == _sha256(second)


def test_gm04_external_relationship_removed_without_text_change(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.docx"
    cleaned = tmp_path / "cleaned.docx"
    _create_docx(source, ("用户可见链接文字",))

    def add_external_relationship(parts: dict[str, bytes]) -> None:
        name = "word/_rels/document.xml.rels"
        root = etree.fromstring(parts[name])
        relation = etree.SubElement(
            root,
            f"{{{_RELATIONSHIP_NAMESPACE}}}Relationship",
        )
        relation.set("Id", "rIdExternal")
        relation.set(
            "Type",
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        )
        relation.set("Target", "https://private.example.invalid/secret")
        relation.set("TargetMode", "External")
        parts[name] = etree.tostring(
            root,
            xml_declaration=True,
            encoding="UTF-8",
        )

    _rewrite_archive(source, add_external_relationship)
    before = [
        element.text
        for element in DocxParser().parse(source, display_path="before.docx")
    ]
    audit = clean_docx(
        source=source,
        destination=cleaned,
        canonical_name="GM-04 质量管理制度及考核办法.docx",
        source_date_epoch=_SOURCE_DATE_EPOCH,
    )
    after = [
        element.text
        for element in DocxParser().parse(cleaned, display_path="after.docx")
    ]

    assert before == after
    assert audit.removed_external_relationship_count == 1
    assert audit.external_relationship_type_counts == {"hyperlink": 1}
    assert b"private.example.invalid" not in cleaned.read_bytes()


@pytest.mark.parametrize(
    ("text", "in_table", "accepted", "level"),
    [
        ("1. 目的", False, True, 1),
        ("2 范围", False, True, 1),
        ("3.1 组织机构及职责", False, True, 2),
        ("3.1.1 质量负责人", False, True, 3),
        ("1. 这是正文。包含完整句子", False, False, None),
        ("1. 表格标题", True, False, None),
        ("普通列表项", False, False, None),
    ],
)
def test_heading_candidate_is_conservative(
    text: str,
    in_table: bool,
    accepted: bool,
    level: int | None,
) -> None:
    decision = heading_candidate(text, in_table=in_table)

    assert decision.accepted is accepted
    assert decision.level == level


def test_heading_repair_preserves_text_order_and_is_idempotent(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.docx"
    first = tmp_path / "first.docx"
    second = tmp_path / "second.docx"
    _create_docx(source, ("1. 目的", "正文保持不变", "1.1 范围"))

    before = [
        element.text
        for element in DocxParser().parse(source, display_path="before.docx")
    ]
    first_audit = clean_docx(
        source=source,
        destination=first,
        canonical_name="GM-02 测试.docx",
        source_date_epoch=_SOURCE_DATE_EPOCH,
    )
    second_audit = clean_docx(
        source=first,
        destination=second,
        canonical_name="GM-02 测试.docx",
        source_date_epoch=_SOURCE_DATE_EPOCH,
    )
    parsed = DocxParser().parse(first, display_path="after.docx")

    assert [element.text for element in parsed] == before
    assert [element.kind for element in parsed] == [
        ElementKind.HEADING,
        ElementKind.PARAGRAPH,
        ElementKind.HEADING,
    ]
    assert first_audit.heading_accepted_count == 2
    assert second_audit.heading_accepted_count == 2
    assert _sha256(first) == _sha256(second)


def test_full_ten_document_corpus_is_exact_private_and_reproducible(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    first = _prepare(
        tmp_path,
        monkeypatch,
        specs=EXPECTED_INVENTORY,
        output_name="first",
    )
    second = _prepare(
        tmp_path,
        monkeypatch,
        specs=EXPECTED_INVENTORY,
        output_name="second",
    )
    expected_names = {
        Path(spec.canonical_name).with_suffix(".docx").name
        for spec in EXPECTED_INVENTORY
    }
    first_docs = {path.name for path in (first / "docs").iterdir()}
    second_docs = {path.name for path in (second / "docs").iterdir()}
    manifest_bytes = (first / "industry-corpus-manifest.json").read_bytes()
    audit = json.loads((first / "industry-corpus-audit.json").read_bytes())

    assert first_docs == second_docs == expected_names
    assert not tuple((first / "reference").iterdir())
    assert not tuple(first.rglob("*.doc"))
    assert all(
        DocxParser().parse(path, display_path=path.name)
        for path in (first / "docs").iterdir()
    )
    assert b"legacy-" not in manifest_bytes
    assert b"/home/" not in manifest_bytes
    assert audit["warnings"] == ["NETWORK_NAMESPACE_UNAVAILABLE"]
    assert {
        path.name: _sha256(path) for path in (first / "docs").iterdir()
    } == {
        path.name: _sha256(path) for path in (second / "docs").iterdir()
    }


def test_failure_is_atomic_and_leaves_no_partial_corpus(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    specs = (
        SourceSpec("GM-01 第一份.doc"),
        SourceSpec("GM-02 第二份.doc"),
    )
    source = tmp_path / "source"
    _write_sources(source, specs)
    converter = _write_fake_soffice(tmp_path / "soffice")
    monkeypatch.setenv("FAKE_FAIL_STEM", "GM-02第二份")
    output = tmp_path / "output"

    with pytest.raises(CorpusPreparationError, match="DOCX_AUDIT_FAILED"):
        prepare_industry_corpus(
            source_dir=source,
            output_root=output,
            libreoffice_path=converter,
            source_date_epoch=_SOURCE_DATE_EPOCH,
            generated_from_git_sha=_GIT_SHA,
            expected_inventory=specs,
        )

    assert output.is_dir()
    assert not tuple(output.iterdir())
